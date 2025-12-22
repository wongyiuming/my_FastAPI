import io
import zipfile
import tarfile
import py7zr
import tempfile
import pathlib
import concurrent.futures
from datetime import datetime
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps
import os

def get_font_path():
    # 优先读取环境变量 (Docker 会传)
    env_path = os.getenv("WATERMARK_FONT_PATH")
    if env_path and pathlib.Path(env_path).exists():
        return env_path

    # Windows 兜底路径
    win_path = "C:\\Windows\\Fonts\\simhei.ttf"
    if pathlib.Path(win_path).exists():
        return win_path

    # 如果都找不到，Pillow 会自动回退到默认字体（虽不支持中文但不会崩）
    return None


def calculate_center(page_w, page_h, text_w, text_h):
    """计算居中坐标的辅助函数"""
    return (page_w - text_w) / 2, (page_h - text_h) / 2


# --- 1. 核心图片处理器 (全屏平铺 + EXIF 修正) ---
def process_single_image(content: bytes, text: str) -> bytes:
    try:
        # 1. 打开图片并修正 EXIF 转向 (解决压缩包图片尺寸偏差的关键)
        raw_img = Image.open(io.BytesIO(content))
        img = ImageOps.exif_transpose(raw_img).convert("RGBA")
        width, height = img.size

        # 2. 准备水印文字
        full_text = f"{text} {datetime.now().strftime('%Y-%m-%d %H:%M')}"

        # 3. 动态计算字号：长边的 1/50
        base_side = max(width, height)
        font_size = int(base_side / 50)
        if font_size < 15: font_size = 15

        try:
            font = ImageFont.truetype(get_font_path(), font_size)
        except:
            font = ImageFont.load_default()

        # 4. 精准计算文字宽高以设定步进
        draw_temp = ImageDraw.Draw(Image.new("RGBA", (1, 1)))
        bbox = draw_temp.textbbox((0, 0), full_text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]

        # 步进：横向留 1.5 倍宽度，纵向留 3 倍高度
        step_x = int(tw * 1.5)
        step_y = int(th * 4)

        # 5. 创建覆盖层 (双倍画布以防旋转留白)
        overlay = Image.new("RGBA", (width * 2, height * 2), (255, 255, 255, 0))
        draw = ImageDraw.Draw(overlay)

        # 6. 平铺绘制 (使用金黄色，透明度加强至 160)
        for x in range(0, width * 2, step_x):
            for y in range(0, height * 2, step_y):
                draw.text((x, y), full_text, fill=(255, 215, 0, 160), font=font)

        # 7. 旋转 45 度并裁剪回原图大小
        overlay = overlay.rotate(45, resample=Image.BICUBIC)
        left = (overlay.width - width) // 2
        top = (overlay.height - height) // 2
        txt_layer = overlay.crop((left, top, left + width, top + height))

        # 8. 合并图层
        combined = Image.alpha_composite(img, txt_layer)
        out = io.BytesIO()
        combined.convert("RGB").save(out, format="JPEG", quality=90)
        return out.getvalue()
    except Exception as e:
        print(f"Image processing failed: {e}")
        return content


# --- 2. PDF 处理器 ---
def process_single_pdf(content: bytes, text: str) -> bytes:
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        mark_text = f"{text} {datetime.now().strftime('%Y-%m-%d')}"
        for page in doc:
            w, h = page.rect.width, page.rect.height
            page.insert_text(
                (w / 4, h / 2), mark_text, fontsize=20,
                color=(1, 0.8, 0), fill_opacity=0.4, rotate=45
            )
        return doc.write()
    except Exception:
        return content


# --- 3. Word 处理器 ---
def process_single_word(content: bytes, text: str) -> bytes:
    try:
        doc = Document(io.BytesIO(content))
        mark_text = f"{text} {datetime.now().strftime('%Y-%m-%d')}"
        for section in doc.sections:
            header = section.header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            run = p.add_run(mark_text)
            run.font.size, run.font.color.rgb = Pt(24), RGBColor(255, 215, 0)
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception:
        return content


# --- 4. 任务调度与过滤 ---
def dispatch_task(item, text):
    name, content = item
    path_obj = pathlib.Path(name)
    ext = path_obj.suffix.lower()

    # 过滤隐藏文件和系统垃圾
    if name.startswith('__MACOSX') or path_obj.name.startswith('._') or ext == '.db':
        return name, content

    try:
        if ext in ['.jpg', '.jpeg', '.png']:
            # 过滤掉极小的缩略图 (小于 5KB)
            if len(content) < 5120: return name, content
            return name, process_single_image(content, text)
        elif ext == '.pdf':
            return name, process_single_pdf(content, text)
        elif ext in ['.docx', '.doc']:
            return name, process_single_word(content, text)
    except Exception as e:
        print(f"Error in dispatcher for {name}: {e}")

    return name, content


def run_batch_task(files_data: list, text: str):
    """利用多线程池进行保序处理"""
    with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
        # map 保证了结果顺序与 files_data 顺序完全一致
        results = list(executor.map(lambda x: dispatch_task(x, text), files_data))
    return results


# --- 5. 通用压缩包处理器 (使用 TemporaryDirectory 以确保兼容性) ---
def process_any_archive(archive_bytes: bytes, text: str, archive_ext: str) -> bytes:
    files_to_work = []
    archive_io = io.BytesIO(archive_bytes)

    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = pathlib.Path(tmp_dir)

            # 分格式解压
            if archive_ext == '.7z':
                with py7zr.SevenZipFile(archive_io, mode='r') as archive:
                    archive.extractall(path=tmp_dir)
            elif archive_ext == '.zip':
                with zipfile.ZipFile(archive_io) as archive:
                    archive.extractall(path=tmp_dir)
            elif '.tar' in archive_ext or archive_ext in ['.tgz', '.tbz2', '.txz']:
                with tarfile.open(fileobj=archive_io, mode='r:*') as archive:
                    archive.extractall(path=tmp_dir)

            # 扫描解压后的物理文件
            for p in tmp_path.rglob('*'):
                if p.is_file():
                    rel_name = str(p.relative_to(tmp_path))
                    files_to_work.append((rel_name, p.read_bytes()))

        # 执行批量水印任务
        processed_list = run_batch_task(files_to_work, text)

        # 重新打包回 ZIP
        out_buffer = io.BytesIO()
        with zipfile.ZipFile(out_buffer, "a", zipfile.ZIP_DEFLATED) as out_zip:
            for name, data in processed_list:
                out_zip.writestr(name, data)
        return out_buffer.getvalue()

    except Exception as e:
        print(f"Archive processing error: {e}")
        return archive_bytes